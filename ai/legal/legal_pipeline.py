# -*- coding: utf-8 -*-
"""사업기획서 → 규제 경로 라우팅(LLM) → 법령 레지스트리 조회 → 법제처 API 조문 페치(결정론)
→ 적합성 선별(LLM) → 리포트.

노트북(aivle_BigProject_v2.ipynb) 대비 변경점
  1. 상태 드리프트 버그 수정: 전 단계를 단일 상태 객체로 연결하고
     "검색계획 경로 ⊆ 라우터 선택 경로"를 assert.
  2. 법령 탐색을 LLM 검색 → 큐레이션 레지스트리(law_registry.json) 조회 + 법제처
     Open API 직접 호출로 교체. LLM은 라우팅(입구)·선별(출구) 두 곳만.

실행:
  python legal_pipeline.py run "기획서.docx" --llm anthropic   # ANTHROPIC_API_KEY 필요 (claude-sonnet-5)
  # 수동(에이전트) 모드 — LLM 호출을 외부에서 수행:
  python legal_pipeline.py prepare-route "기획서.docx"          # → 출력/작업/route_prompt.txt
  python legal_pipeline.py apply-route 출력/작업/route_result.json
  python legal_pipeline.py apply-screen 출력/작업/screen_result.json
"""
import argparse
import json
import os
import re
import shutil
import sys
import time
import urllib.parse
import urllib.request
from datetime import date
from pathlib import Path

BASE_DIR = Path(__file__).resolve().parent
REGISTRY_PATH = BASE_DIR / "law_registry.json"
CATEGORY_MAP_PATH = BASE_DIR / "category_map.json"
OUT_DIR = BASE_DIR / "출력"
WORK_DIR = OUT_DIR / "작업"
# 법령 캐시는 작업 디렉터리와 무관하게 항상 공유한다 (읽기 중심, 재사용 이득이 크다).
CACHE_DIR = WORK_DIR / "law_cache"

# 법제처 Open API OC(이메일 ID). 기본값 test는 공개 테스트 ID이므로 운영 전 발급 권장.
LAW_API_OC = os.getenv("LAW_API_OC") or "test"
LAW_SEARCH_URL = "http://www.law.go.kr/DRF/lawSearch.do"
LAW_SERVICE_URL = "http://www.law.go.kr/DRF/lawService.do"

MODEL = os.getenv("LEGAL_MODEL") or "claude-sonnet-5"


class PipelineError(Exception):
    """파이프라인이 입력·상태 문제로 진행할 수 없을 때. 재시도해도 소용없는 오류."""


class BackendUnavailableError(PipelineError):
    """LLM·외부 API 호출이 실패했을 때. 입력 문제가 아니므로 재시도할 가치가 있다."""

ROUTE_STATUSES = ("해당", "적용 가능", "비해당", "불명")
SCREEN_CATEGORIES = ("requirement", "risk", "scope", "supporting", "exclude")

# ---------------------------------------------------------------- 프롬프트

ROUTING_INSTRUCTIONS = """당신은 상품기획서에서 규제 조사 경로를 선택하는 범용 규제 라우터다.
법률의 최종 적용 여부를 판단하지 않는다. 구체적인 법령명·조문·URL을 생성하지 않는다.
제공된 기획서 원문만 사용한다.

[핵심 원칙]
1. 상품명을 보고 규제 분야를 추측하지 않는다.
2. 모든 '해당'/'적용 가능' 판단에는 기획서 원문 인용을 evidence_quotes에 그대로 복사해 넣는다.
3. 원문 근거가 없으면 '불명'으로 판단한다.
4. 소재가 존재한다는 이유만으로 소비재 안전 규제를 자동 선택하지 않는다. 다만 일반 소비자가
   사용하는 실물 제품이면 consumer_product_safety는 최소 '적용 가능'으로 조사 대상에 넣는다.
5. 관련 단어가 등장하는 것과 규제 대상 행위를 수행하는 것을 구분한다.
6. 목록에 없는 규제 경로는 additional_route_candidates에 작성한다.

[적용범위 규칙]
- food_contact: 사람이 섭취할 식품·음료와 직접 접촉하도록 설계된 경우만. 폐기물·음식물쓰레기 등
  폐기 대상 물질과만 접촉하는 제품은 선택하지 않는다.
- environmental_waste: 사업자가 폐기물 수집·운반·처리·재활용 등 환경규제 대상 행위를 직접
  수행하는 경우만. 단순히 폐기물을 담는 상품이라는 이유만으로 선택하지 않는다.
- chemical_biocidal: 제품 자체가 세정·살균·소독·방충 등 화학적 기능을 주목적으로 하는 경우만.
- 라우팅 단계에서는 등록·신고·인증·허가 의무를 단정하지 않는다.

[규제 경로 정의]
consumer_product_safety: 일반 소비자용 실물 상품의 제품안전·표시 관련
electrical_safety: 전원·배터리·충전·전기부품 사용 상품
wireless_radio: Wi-Fi, Bluetooth, NFC 등 무선 송수신 기능
food_product: 사람이 섭취하는 식품·음료
food_contact: 식품과 직접 접촉하는 기구·용기·포장·조리도구
health_functional_food: 건강기능 효과를 표시하는 섭취 제품
medical_device: 질병 진단·치료·예방 등 의료 목적 기기
pharmaceutical: 의약품·의약외품
cosmetic: 피부·모발 등 청결·미화 목적 상품
chemical_biocidal: 세정제·살균제·소독제·살충제 등 생활화학·살생물 제품
children_product: 어린이·유아를 주요 사용자로 설계한 상품
industrial_safety: 산업현장·작업자 안전·보호구 관련
vehicle_mobility: 자동차·이동수단·차량 부품
construction_building: 건축자재·건설설비
environmental_waste: 폐기물 처리·재활용 등 환경규제 행위를 직접 수행
personal_data: 개인정보 수집·저장·분석
location_data: 위치정보 수집·이용
online_sales: 자사몰·온라인몰·오픈마켓·펀딩 등 비대면 판매
advertising_claims: 제거율·감소율 등 검증 가능한 성능 광고 표현
import_export_customs: 수입·수출·해외 제조
intellectual_property: 특허·상표·디자인권·저작권 검토가 명시된 경우
finance_payment: 금융·대출·보험·결제·전자금융 기능
education_service: 학습·교육·강의 서비스
telecom_platform: 통신서비스·중개 플랫폼 운영
software_ai_service: 소프트웨어·SaaS·AI 분석 서비스 제공
content_copyright: 영상·음악·이미지 등 저작권 콘텐츠 제작·배포
other: 위에 없지만 별도 규제 가능성이 있는 경우

[상태 기준]
해당: 기획서에서 해당 기능·사업 행위가 명확히 확인됨
적용 가능: 관련 가능성은 있으나 제품 분류·용도·세부조건 확인 필요
비해당: 원문에서 명확히 부정되거나 정의상 직접 배제됨
불명: 판단에 필요한 정보가 없음

[출력]
반드시 아래 스키마의 JSON 하나만 출력한다. 코드블록·설명 금지.
routes에는 '해당'/'적용 가능'인 경로와, 중요한 '불명' 경로만 담는다(무관 경로 나열 금지).
{
  "product_facts": {
    "product_name": str, "product_type": str, "intended_use": str,
    "materials": [str], "core_functions": [str], "target_users": [str],
    "sales_channels": [str], "performance_claims": [str], "personal_data_features": [str]
  },
  "routes": [
    {"route_id": str, "status": "해당"|"적용 가능"|"비해당"|"불명",
     "evidence_quotes": [str], "reason": str, "confidence": float}
  ],
  "additional_route_candidates": [{"route_name": str, "reason": str}],
  "missing_information": [
    {"question": str, "related_route_ids": [str]}
  ]
}
missing_information의 related_route_ids에는 그 질문의 답이 판단에 영향을 주는
규제 경로 id를 담는다(위 [규제 경로 정의] 목록의 id만 사용)."""

REVISION_INSTRUCTIONS = """입력은 상품기획서 원문과 규제 검토 결과(직접 의무·제재 조문, 권장 조치)다.
기획서 '문장 자체'가 규제 위반 소지를 만들어 고쳐 써야 하는 경우에만 수정 요청을 만든다.

[규칙]
1. quote에는 기획서 원문의 해당 문장을 한 글자도 바꾸지 말고 그대로 복사한다.
   [확정 정보] 블록의 문장은 기획서 본문이 아니므로 인용하지 않는다.
2. 확인 질문으로 해결될 정보 부족 문제는 수정 요청으로 만들지 않는다.
3. 각 요청에는 서로 다른 방향의 수정안을 2~3개 제시한다 (label A/B/C).
   new_text는 quote를 대체할 완결된 문장이어야 한다.
4. 고칠 문장이 없으면 빈 배열을 출력한다. 억지로 만들지 않는다.

[출력]
반드시 아래 스키마의 JSON 하나만 출력한다. 코드블록·설명 금지.
{
  "revision_requests": [
    {"route_id": str, "quote": str, "rationale": str,
     "suggestions": [{"label": "A"|"B"|"C", "new_text": str}]}
  ]
}"""

SCREENING_INSTRUCTIONS = """입력은 상품 규제 프로필과, 큐레이션 레지스트리·법제처 API로 확보한 현행 법령 조문 후보다.
각 후보가 이 상품의 규제 조사에서 어떤 용도인지 분류한다. 법령의 실존·현행성은 이미 API로
확인됐으므로 판단하지 않는다. 최종 적용 여부도 확정하지 않는다.

[분류 규칙]
1. 모든 citation_id에 대해 평가를 하나씩 반환한다. 없는 id를 만들지 않는다.
2. category:
   - requirement: 사업자·제조자·판매자에게 직접 의무·금지를 부과하고 이 상품에 관련
   - risk: 과태료·벌칙·손해배상 등 위반 시 제재 조항
   - scope: 적용 대상·품목·범위를 판단하는 조항 (이 상품의 해당 여부 확인에 필요)
   - supporting: 정의·해석 보조 조항
   - exclude: 이 상품·판매방식·광고표현과 무관
3. 국가·행정기관에만 의무를 부과하는 조항은 requirement가 될 수 없다.
4. relevance_note에는 이 상품 기준의 관련 이유를 1문장으로 쓴다.
5. plain_summary에는 그 조문 자체가 요구·금지하는 내용을 법률 용어 없이 1~2문장으로 쓴다.
   - 상품과의 관련성이 아니라 조문의 내용을 설명한다 (관련성은 relevance_note 몫이다).
   - "무엇을 해야 하는가/하면 안 되는가"를 행동으로 쓴다.
     예: "온라인으로 물건을 팔려면 관할 시·군·구청에 통신판매업 신고를 해야 합니다."
   - 조문에 없는 의무를 만들지 않는다. 조문이 정의·절차뿐이면 그대로 그렇게 쓴다.
   - category가 exclude인 후보는 빈 문자열로 둔다.

[결론 도출]
분류를 마친 뒤, requirement/scope 후보를 근거로 "이 사업자가 지금(판매 개시 전) 실제로
해야 할 일"을 위험도 순으로 최대 5개 뽑는다(action_items).
- 하나의 action은 여러 조문을 묶은 실무 행동 1개다 (예: "통신판매업 신고 + 신원 표시").
- 미래 계획(출원·수출 등)에만 해당하는 조건부 의무는 Top 5에서 제외하고 timing이
  "계획 실행 시"인 항목으로만 최대 2개 별도로 담는다.
- 벌칙 조항은 action이 아니라 해당 action의 reason에 제재 수위로 언급한다.

[출력]
반드시 아래 스키마의 JSON 하나만 출력한다. 코드블록·설명 금지.
{
  "screenings": [
    {"citation_id": str, "category": "requirement"|"risk"|"scope"|"supporting"|"exclude",
     "relevance_note": str, "plain_summary": str}
  ],
  "route_gaps": [str],  // 직접 의무 근거(requirement)가 하나도 없는 route_id 목록
  "action_items": [
    {"rank": int, "action": str, "reason": str, "citation_ids": [str],
     "timing": "즉시"|"판매 개시 전"|"계획 실행 시"}
  ]
}"""

# ---------------------------------------------------------------- 유틸

def log(msg):
    print(msg, flush=True)


def normalize(text):
    return re.sub(r"\s+", "", text or "")


def read_docx(path):
    from docx import Document
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(dict.fromkeys(cells)))
    return "\n".join(parts)


def build_source_text(sections, confirmed_facts=None):
    """확정된 StructuredPlan 섹션 목록 → 라우팅 원문. docx 대신 쓰는 입력 경로.

    confirmed_facts(질문 답변으로 확정된 정보)는 반드시 이 산출물 자체에 붙인다.
    validate_routing이 인용을 source_text 부분문자열 대조로 화이트리스트하므로,
    별도 프롬프트 블록에 넣으면 확정 정보를 인용한 근거가 전부 삭제된다.
    """
    parts = []
    for s in sections or []:
        content = (s.get("content") or "").strip()
        if not content:
            continue
        title = (s.get("title") or s.get("code") or "").strip()
        parts.append(f"## {title}\n{content}" if title else content)
    if not parts:
        raise PipelineError("확정된 사업계획에 조사할 본문이 없습니다.")
    text = "\n\n".join(parts)

    facts = [f for f in (confirmed_facts or []) if f.get("key") and f.get("value")]
    if facts:
        lines = ["[확정 정보 — 아래 사실은 사용자가 외부 확인을 마친 것으로 판단에 그대로 사용하라]"]
        for f in facts:
            line = f"- {f['key']}: {f['value']}"
            meta = ", ".join(x for x in (f.get("source"), (f.get("answeredAt") or "")[:10]) if x)
            if meta:
                line += f" (출처: {meta})"
            lines.append(line)
        text = text + "\n\n" + "\n".join(lines)
        log(f"확정 정보 {len(facts)}건 주입: {[f['key'] for f in facts]}")
    return text


def load_registry():
    return json.loads(REGISTRY_PATH.read_text(encoding="utf-8"))


def parse_llm_json(text):
    text = text.strip()
    if text.startswith("```"):
        text = re.sub(r"^```[a-zA-Z]*\n?|\n?```$", "", text).strip()
    start, end = text.find("{"), text.rfind("}")
    if start == -1 or end == -1:
        raise ValueError("LLM 응답에서 JSON을 찾지 못했습니다.")
    return json.loads(text[start:end + 1])


# ---------------------------------------------------------------- 검증 (라우팅)

def validate_routing(raw, source_text):
    """라우팅 JSON 스키마·근거 검증. 원문에 없는 인용 제거, 근거 소실 시 '해당'→'적용 가능' 강등."""
    audit = {"removed_quotes": [], "downgraded": [], "dropped_routes": []}
    norm_src = normalize(source_text)
    facts = raw.get("product_facts") or {}
    routes = []
    seen = set()
    for r in raw.get("routes") or []:
        rid, status = r.get("route_id"), r.get("status")
        if not rid or status not in ROUTE_STATUSES or rid in seen:
            audit["dropped_routes"].append(rid)
            continue
        seen.add(rid)
        quotes = []
        for q in r.get("evidence_quotes") or []:
            if normalize(q) and normalize(q) in norm_src:
                quotes.append(q)
            else:
                audit["removed_quotes"].append({"route_id": rid, "quote": q})
        if status == "해당" and not quotes:
            status = "적용 가능"
            audit["downgraded"].append(rid)
        routes.append({
            "route_id": rid, "status": status, "evidence_quotes": quotes,
            "reason": r.get("reason", ""), "confidence": r.get("confidence"),
        })
    return {
        "product_facts": facts,
        "routes": routes,
        "additional_route_candidates": raw.get("additional_route_candidates") or [],
        "missing_information": _normalize_missing_information(raw.get("missing_information")),
        "audit": audit,
    }


def _normalize_missing_information(items):
    """missing_information을 {question, related_route_ids} 형태로 정규화.

    구 스키마(평문 문자열 배열)와 구 state 파일도 그대로 수용한다.
    """
    normalized = []
    for item in items or []:
        if isinstance(item, str) and item.strip():
            normalized.append({"question": item.strip(), "related_route_ids": []})
        elif isinstance(item, dict) and (item.get("question") or "").strip():
            route_ids = [r for r in (item.get("related_route_ids") or []) if isinstance(r, str)]
            normalized.append({"question": item["question"].strip(),
                               "related_route_ids": route_ids})
    return normalized


# ---------------------------------------------------------------- 법제처 API

def _api_get(url, params):
    qs = urllib.parse.urlencode(params, encoding="utf-8")
    req = urllib.request.Request(f"{url}?{qs}", headers={"User-Agent": "legal-pipeline/1.0"})
    with urllib.request.urlopen(req, timeout=30) as resp:
        return json.loads(resp.read().decode("utf-8"))


def search_law_exact(name):
    """법령명 정확 일치 + 현행 법령 1건의 검색 메타 반환 (없으면 None)."""
    data = _api_get(LAW_SEARCH_URL, {
        "OC": LAW_API_OC, "target": "law", "type": "JSON", "query": name, "display": 50,
    })
    items = data.get("LawSearch", {}).get("law") or []
    if isinstance(items, dict):
        items = [items]
    for item in items:
        if item.get("법령명한글", "").replace(" ", "") == name.replace(" ", "") \
                and item.get("현행연혁코드") == "현행":
            return item
    return None


def fetch_law_articles(mst):
    """MST로 현행 법령 전문을 받아 조문 리스트로 정규화. law_cache에 캐시."""
    CACHE_DIR.mkdir(parents=True, exist_ok=True)
    cache = CACHE_DIR / f"{mst}.json"
    if cache.exists():
        data = json.loads(cache.read_text(encoding="utf-8"))
    else:
        data = _api_get(LAW_SERVICE_URL, {"OC": LAW_API_OC, "target": "law", "MST": mst, "type": "JSON"})
        cache.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")
        time.sleep(0.3)
    units = data.get("법령", {}).get("조문", {}).get("조문단위") or []
    if isinstance(units, dict):
        units = [units]
    articles = []
    for u in units:
        if u.get("조문여부") != "조문":
            continue
        body = [u.get("조문내용", "")]
        hang = u.get("항") or []
        if isinstance(hang, dict):
            hang = [hang]
        for h in hang:
            body.append(h.get("항내용", ""))
            ho = h.get("호") or []
            if isinstance(ho, dict):
                ho = [ho]
            for x in ho:
                body.append(x.get("호내용", ""))
        text = "\n".join(t.strip() for t in body if isinstance(t, str) and t.strip())
        articles.append({
            "조문번호": str(u.get("조문번호", "")),
            "조문가지번호": str(u.get("조문가지번호", "") or ""),
            "제목": u.get("조문제목", "") or "",
            "내용": text,
        })
    return articles


def filter_articles(articles, keywords, min_keep=3, max_keep=25, excerpt_len=350):
    """focus_keywords로 조문을 거른다. 결과가 너무 적으면 전체(축약)로 폴백.

    제목이 키워드와 맞는 조문을 먼저 남긴다. 조문번호 순으로만 자르면
    앞 순번 조문이 자리를 다 차지해, 정작 제목이 정확히 맞는 뒤 순번 조문
    (예: 전안법 제28조 안전기준준수대상, 개인정보법 제30조 처리방침)이
    max_keep 밖으로 밀려나 구조적으로 도달할 수 없게 된다.
    """
    def rank(a):
        title = a["제목"] or ""
        if any(k in title for k in keywords):
            return 0
        if any(k in a["내용"][:200] for k in keywords):
            return 1
        return 2

    if keywords:
        matched = [(rank(a), i, a) for i, a in enumerate(articles)]
        matched = [m for m in matched if m[0] < 2]
        # 제목 일치를 먼저 취해 max_keep 안에 넣고, 남길 것을 정한 뒤
        # 다시 조문 순서로 되돌린다 (리포트 가독성).
        kept = sorted(sorted(matched, key=lambda m: (m[0], m[1]))[:max_keep],
                      key=lambda m: m[1])
        picked = [a for _, _, a in kept]
    else:
        picked = list(articles)
    fallback = False
    if len(picked) < min_keep:
        picked, fallback = list(articles), True
    picked = picked[:max_keep]
    for a in picked:
        if len(a["내용"]) > excerpt_len:
            a["내용"] = a["내용"][:excerpt_len] + " …(생략)"
    return picked, fallback


# ---------------------------------------------------------------- 증분 재검토

def load_category_map():
    return json.loads(CATEGORY_MAP_PATH.read_text(encoding="utf-8"))


def filter_routes_for_categories(profile, rerun_categories, category_map):
    """증분 재검토: 재실행 범주에 매핑되지 않는 '해당/적용 가능' 라우트를 계획에서 제외한다.

    라우트는 category_map의 매핑 범주가 rerun_categories와 하나라도 겹치면 유지된다.
    비해당/불명 라우트는 수집 대상이 아니므로 그대로 둔다.
    """
    rerun = set(rerun_categories or [])
    if not rerun:
        return profile
    route_map = category_map["routes"]
    kept, dropped = [], []
    for route in profile["routes"]:
        if route["status"] not in ("해당", "적용 가능"):
            kept.append(route)
            continue
        mapped = set(route_map.get(route["route_id"]) or ["INDUSTRY_SPECIFIC"])
        if mapped & rerun:
            kept.append(route)
        else:
            dropped.append(route["route_id"])
    kept_active = [r["route_id"] for r in kept if r["status"] in ("해당", "적용 가능")]
    log(f"증분 재검토: 실행 경로 {kept_active} / 제외 경로 {dropped} "
        f"(유지 범주는 재실행하지 않음, rerun={sorted(rerun)})")
    filtered = dict(profile)
    filtered["routes"] = kept
    return filtered


# ---------------------------------------------------------------- 계획·수집

def build_plan(profile, registry):
    """라우터 결과 → 레지스트리 조회 계획. (버그1 수정: 선택 경로와 계획 경로 불일치 차단)"""
    selected = {r["route_id"]: r for r in profile["routes"] if r["status"] in ("해당", "적용 가능")}
    routes_cfg = registry["routes"]
    plan, unknown = [], []
    for rid, r in selected.items():
        cfg = routes_cfg.get(rid)
        if cfg is None:
            unknown.append(rid)
            continue
        plan.append({
            "route_id": rid, "status": r["status"], "topic": cfg["topic"],
            "priority": cfg["priority"], "laws": cfg["laws"], "notices": cfg.get("notices", []),
        })
    plan.sort(key=lambda p: (p["status"] != "해당", -p["priority"]))
    plan_ids = {p["route_id"] for p in plan}
    if not plan_ids <= set(selected):  # 상태 드리프트 방지
        raise PipelineError(f"계획 경로가 라우터 선택을 벗어남: {plan_ids - set(selected)}")
    return plan, unknown


def collect_candidates(plan):
    """계획의 각 법령을 법제처 API로 페치해 CIT 후보를 만든다 (LLM 불개입)."""
    candidates, fetch_log = [], []
    n = 0
    for p in plan:
        for law in p["laws"]:
            name = law["name"]
            meta = search_law_exact(name)
            if meta is None:
                fetch_log.append({"law": name, "route": p["route_id"], "status": "not_found"})
                log(f"  [실패] {name}: 현행 법령 정확 일치 없음")
                continue
            articles = fetch_law_articles(meta["법령일련번호"])
            picked, fallback = filter_articles(articles, law.get("focus_keywords") or [])
            fetch_log.append({
                "law": name, "route": p["route_id"], "status": "ok",
                "mst": meta["법령일련번호"], "law_id": meta.get("법령ID"),
                "시행일자": meta.get("시행일자"), "total_articles": len(articles),
                "picked": len(picked), "keyword_fallback": fallback,
            })
            log(f"  [수집] {p['route_id']} / {name} (MST {meta['법령일련번호']}, "
                f"시행 {meta.get('시행일자')}) 조문 {len(articles)} → 후보 {len(picked)}")
            for a in picked:
                n += 1
                jo = f"제{a['조문번호']}조" + (f"의{a['조문가지번호']}" if a["조문가지번호"] else "")
                candidates.append({
                    "citation_id": f"CIT-{n:03d}", "route_id": p["route_id"], "law_name": name,
                    "조문": jo, "제목": a["제목"], "내용": a["내용"],
                    "mst": meta["법령일련번호"], "현행": True, "시행일자": meta.get("시행일자"),
                })
    return candidates, fetch_log


# ---------------------------------------------------------------- 검증 (선별)

def validate_screening(raw, candidates):
    valid_ids = {c["citation_id"] for c in candidates}
    result, seen = {}, set()
    invented = []
    for s in raw.get("screenings") or []:
        cid = s.get("citation_id")
        if cid not in valid_ids:
            invented.append(cid)
            continue
        cat = s.get("category")
        if cat not in SCREEN_CATEGORIES:
            cat = "exclude"
        # plain_summary는 후행 도입 필드다 — 없는 응답(구 프롬프트로 저장된 work_dir)도 그대로 받는다.
        result[cid] = {"category": cat, "relevance_note": s.get("relevance_note", ""),
                       "plain_summary": (s.get("plain_summary") or "").strip()}
        seen.add(cid)
    missing = sorted(valid_ids - seen)
    for cid in missing:
        result[cid] = {"category": "exclude", "relevance_note": "(LLM 미평가 — 자동 제외)",
                       "plain_summary": ""}
    actions = []
    for a in raw.get("action_items") or []:
        if not a.get("action"):
            continue
        actions.append({
            "rank": a.get("rank", len(actions) + 1),
            "action": a["action"],
            "reason": a.get("reason", ""),
            "citation_ids": [c for c in (a.get("citation_ids") or []) if c in valid_ids],
            "timing": a.get("timing", ""),
        })
    actions.sort(key=lambda a: ((a["timing"] == "계획 실행 시"), a["rank"]))
    return result, {"invented_ids": invented, "missing_ids": missing,
                    "route_gaps": raw.get("route_gaps") or [],
                    "action_items": actions}


# ---------------------------------------------------------------- LLM 백엔드

_RETRY_NAG = "\n\n(이전 응답이 유효한 JSON이 아니었다. 스키마에 맞는 JSON 하나만 출력하라.)"


def call_anthropic(prompt, system):
    """Anthropic SDK 직접 호출. ANTHROPIC_API_KEY 필요."""
    import anthropic
    client = anthropic.Anthropic()
    for attempt in (1, 2):
        resp = client.messages.create(
            model=MODEL, max_tokens=16000, system=system,
            messages=[{"role": "user", "content": prompt}],
        )
        text = "".join(b.text for b in resp.content if b.type == "text")
        try:
            return parse_llm_json(text)
        except (ValueError, json.JSONDecodeError):
            if attempt == 2:
                raise
            prompt = prompt + _RETRY_NAG


def call_claude_cli(prompt, system):
    """Claude Code CLI 헤드리스 호출. API 키 대신 CLI 로그인 세션을 쓴다.

    호출마다 Claude Code 자체 시스템 프롬프트(~47k 토큰)가 함께 청구되므로
    API 키를 쓸 수 있으면 그쪽이 싸다. 로컬 데모용 경로다.
    """
    import subprocess
    timeout = int(os.getenv("LEGAL_CLI_TIMEOUT") or "900")
    for attempt in (1, 2):
        proc = subprocess.run(
            ["claude", "-p", "--model", MODEL,
             "--append-system-prompt", system,
             "--allowed-tools", "",          # 파일시스템·툴 접근 차단
             "--output-format", "json"],
            input=prompt, capture_output=True, text=True,
            encoding="utf-8", timeout=timeout,
        )
        if proc.returncode != 0:
            raise BackendUnavailableError(
                f"claude CLI 실패(exit {proc.returncode}): {(proc.stderr or '').strip()[:300]}")
        try:
            envelope = json.loads(proc.stdout)
        except json.JSONDecodeError:
            raise BackendUnavailableError("claude CLI 응답을 JSON으로 읽지 못했습니다.")
        if envelope.get("is_error"):
            raise BackendUnavailableError(f"claude CLI 오류: {str(envelope.get('result'))[:300]}")
        try:
            return parse_llm_json(envelope.get("result") or "")
        except (ValueError, json.JSONDecodeError):
            if attempt == 2:
                raise
            prompt = prompt + _RETRY_NAG


def resolve_llm_backend(name=None):
    """사용할 LLM 백엔드를 고른다. 지정이 없으면 쓸 수 있는 쪽을 자동 선택한다."""
    name = (name or os.getenv("LEGAL_LLM_BACKEND") or "auto").strip().lower()
    if name == "auto":
        if os.getenv("ANTHROPIC_API_KEY"):
            return "anthropic"
        if shutil.which("claude"):
            return "claude-cli"
        raise PipelineError(
            "LLM 백엔드가 없습니다. ANTHROPIC_API_KEY를 설정하거나 claude CLI에 로그인하세요.")
    if name not in ("anthropic", "claude-cli"):
        raise PipelineError(f"알 수 없는 LLM 백엔드: {name}")
    return name


def call_llm(prompt, system, backend=None):
    backend = resolve_llm_backend(backend)
    return call_anthropic(prompt, system) if backend == "anthropic" \
        else call_claude_cli(prompt, system)


# ---------------------------------------------------------------- 상태·리포트

def save_state(state, work_dir=None):
    work_dir = Path(work_dir or WORK_DIR)
    work_dir.mkdir(parents=True, exist_ok=True)
    (work_dir / "state.json").write_text(
        json.dumps(state, ensure_ascii=False, indent=2), encoding="utf-8")


def load_state(work_dir=None):
    p = Path(work_dir or WORK_DIR) / "state.json"
    if not p.exists():
        raise PipelineError("state.json이 없습니다. prepare-route부터 실행하세요.")
    return json.loads(p.read_text(encoding="utf-8"))


def write_report(state, screenings, screen_audit, out_dir=None):
    profile, plan = state["profile"], state["plan"]
    candidates = state["candidates"]
    by_cat = {c: [] for c in SCREEN_CATEGORIES}
    for c in candidates:
        s = screenings[c["citation_id"]]
        by_cat[s["category"]].append((c, s))

    facts = profile["product_facts"]
    by_id = {c["citation_id"]: c for c in candidates}
    lines = [
        f"# 법령 조사 리포트 — {facts.get('product_name', '(상품명 미확인)')}",
        f"- 생성일: {date.today().isoformat()}  |  라우팅·선별 모델: {state.get('llm', MODEL)}",
        f"- 원본 기획서: {state.get('source_label') or state.get('docx', '(미지정)')}",
        "",
        "## 결론 — 지금 해야 할 일",
    ]
    actions = screen_audit.get("action_items") or []
    now_items = [a for a in actions if a["timing"] != "계획 실행 시"]
    later_items = [a for a in actions if a["timing"] == "계획 실행 시"]
    if not actions:
        lines.append("- (LLM이 action_items를 반환하지 않음 — 4장 requirement에서 직접 도출 필요)")
    for a in now_items:
        laws = "; ".join(dict.fromkeys(
            f"{by_id[c]['law_name']} {by_id[c]['조문']}" for c in a["citation_ids"] if c in by_id))
        lines.append(f"{a['rank']}. **{a['action']}** ({a['timing']}) — {a['reason']}")
        if laws:
            lines.append(f"   - 근거: {laws}")
    if later_items:
        lines.append("")
        lines.append("계획 실행 시 (지금은 대상 아님):")
        for a in later_items:
            laws = "; ".join(dict.fromkeys(
                f"{by_id[c]['law_name']} {by_id[c]['조문']}" for c in a["citation_ids"] if c in by_id))
            lines.append(f"- **{a['action']}** — {a['reason']}" + (f" (근거: {laws})" if laws else ""))
    lines += [
        "",
        "## 1. 상품 요약",
        f"- 유형: {facts.get('product_type', '')} / 용도: {facts.get('intended_use', '')}",
        f"- 판매채널: {', '.join(facts.get('sales_channels') or [])}",
        f"- 성능 주장: {', '.join(facts.get('performance_claims') or [])}",
        "",
        "## 2. 선택된 규제 경로",
    ]
    for r in profile["routes"]:
        if r["status"] in ("해당", "적용 가능"):
            lines.append(f"- **{r['route_id']}** ({r['status']}) — {r['reason']}")
    lines += ["", "## 3. 조사 대상 법령 (레지스트리 조회 + 법제처 API 현행 확인)"]
    for f in state["fetch_log"]:
        if f["status"] == "ok":
            lines.append(f"- [{f['route']}] {f['law']} — MST {f['mst']}, 시행일 {f['시행일자']}, "
                         f"조문 {f['total_articles']}건 중 {f['picked']}건 후보화")
        else:
            lines.append(f"- [{f['route']}] {f['law']} — **조회 실패({f['status']})**")

    def section(title, cat):
        lines.append("")
        lines.append(f"## {title}")
        rows = by_cat[cat]
        if not rows:
            lines.append("- (없음)")
        for c, s in rows:
            lines.append(f"- **[{c['citation_id']}] {c['law_name']} {c['조문']}({c['제목']})** — {s['relevance_note']}")

    section("4. 직접 의무·금지 (requirement)", "requirement")
    section("5. 적용범위 판단 (scope)", "scope")
    section("6. 위반 시 제재 (risk)", "risk")
    section("7. 보조·정의 (supporting)", "supporting")

    lines += ["", "## 8. 고시·행정규칙 확인 필요 (API 미조회 — 담당자 확인)"]
    notices = [(p["route_id"], n) for p in plan for n in p["notices"]]
    lines += [f"- [{rid}] {n}" for rid, n in notices] or ["- (없음)"]

    lines += ["", "## 9. 검증 로그"]
    lines.append(f"- 라우팅 근거 검증: 인용 제거 {len(profile['audit']['removed_quotes'])}건, "
                 f"강등 {profile['audit']['downgraded']}")
    lines.append(f"- 선별 검증: 미평가 {len(screen_audit['missing_ids'])}건, "
                 f"무효 id {len(screen_audit['invented_ids'])}건")
    gaps = screen_audit.get("route_gaps") or []
    lines.append(f"- requirement 공백 경로: {gaps if gaps else '없음'}")
    excl = len(by_cat["exclude"])
    lines.append(f"- 제외 후보: {excl}건 / 전체 {len(candidates)}건")
    if state.get("unknown_routes"):
        lines.append(f"- 레지스트리 미등록 경로(수동 확인 필요): {state['unknown_routes']}")

    out_dir = Path(out_dir or OUT_DIR)
    out_dir.mkdir(parents=True, exist_ok=True)
    path = out_dir / "법령조사_리포트.md"
    path.write_text("\n".join(lines), encoding="utf-8")
    return path


# ---------------------------------------------------------------- 단계

def stage_prepare(docx_path, work_dir=None):
    text = read_docx(docx_path)
    log(f"기획서 로드: {docx_path} ({len(text):,}자)")
    return stage_prepare_text(text, str(docx_path), work_dir)


def stage_prepare_text(text, source_label, work_dir=None):
    """docx 여부와 무관하게 원문 텍스트로 라우팅 프롬프트와 초기 상태를 만든다."""
    work_dir = Path(work_dir or WORK_DIR)
    prompt = f"[기획서 원문]\n{text}"
    work_dir.mkdir(parents=True, exist_ok=True)
    (work_dir / "route_prompt.txt").write_text(
        ROUTING_INSTRUCTIONS + "\n\n" + prompt, encoding="utf-8")
    save_state({"docx": source_label, "source_label": source_label, "source_text": text}, work_dir)
    log(f"라우팅 프롬프트 저장: {work_dir / 'route_prompt.txt'}")
    return text, prompt


def stage_route(state, routing_raw, work_dir=None, rerun_categories=None):
    profile = validate_routing(routing_raw, state["source_text"])
    if rerun_categories:
        profile = filter_routes_for_categories(profile, rerun_categories, load_category_map())
    log("선택된 경로: " + ", ".join(
        f"{r['route_id']}({r['status']})" for r in profile["routes"]
        if r["status"] in ("해당", "적용 가능")))
    if profile["audit"]["downgraded"] or profile["audit"]["removed_quotes"]:
        log(f"근거 검증: 강등 {profile['audit']['downgraded']}, "
            f"인용 제거 {len(profile['audit']['removed_quotes'])}건")
    registry = load_registry()
    plan, unknown = build_plan(profile, registry)
    log("법령 수집 시작 (법제처 Open API)...")
    candidates, fetch_log = collect_candidates(plan)
    log(f"후보 조문 총 {len(candidates)}건")

    state.update({"profile": profile, "plan": plan, "candidates": candidates,
                  "fetch_log": fetch_log, "unknown_routes": unknown})
    # 선별 프롬프트
    routes_txt = "\n".join(f"- {r['route_id']} ({r['status']}): {r['reason']}"
                           for r in profile["routes"] if r["status"] in ("해당", "적용 가능"))
    cand_txt = "\n\n".join(
        f"[{c['citation_id']}] ({c['route_id']}) {c['law_name']} {c['조문']} {c['제목']}\n{c['내용']}"
        for c in candidates)
    prompt = (f"[상품 프로필]\n{json.dumps(state['profile']['product_facts'], ensure_ascii=False, indent=1)}\n\n"
              f"[선택된 규제 경로]\n{routes_txt}\n\n[법령 조문 후보]\n{cand_txt}")
    work_dir = Path(work_dir or WORK_DIR)
    work_dir.mkdir(parents=True, exist_ok=True)
    (work_dir / "screen_prompt.txt").write_text(
        SCREENING_INSTRUCTIONS + "\n\n" + prompt, encoding="utf-8")
    save_state(state, work_dir)
    log(f"선별 프롬프트 저장: {work_dir / 'screen_prompt.txt'}")
    return state, prompt


def stage_screen(state, screening_raw, out_dir=None):
    screenings, audit = validate_screening(screening_raw, state["candidates"])
    path = write_report(state, screenings, audit, out_dir)
    log(f"리포트 저장: {path}")
    return path


# ---------------------------------------------------------------- 수정안 생성

def build_revision_prompt(state, screenings, screen_audit):
    """수정안 프롬프트: 원문 + 근거로 채택된 조문(requirement/risk) + 권장 조치."""
    candidates = state.get("candidates") or []
    relevant = []
    for c in candidates:
        s = screenings.get(c["citation_id"]) or {}
        if s.get("category") in ("requirement", "risk"):
            relevant.append(f"[{c['citation_id']}] ({c['route_id']}) {c['law_name']} "
                            f"{c['조문']} {c['제목']} — {s.get('relevance_note', '')}")
    actions = (screen_audit or {}).get("action_items") or []
    action_txt = "\n".join(f"- {a['action']} ({a.get('timing') or '시점 미정'}): {a.get('reason', '')}"
                           for a in actions)
    return (f"[기획서 원문]\n{state['source_text']}\n\n"
            f"[직접 의무·제재 조문]\n" + ("\n".join(relevant) or "(없음)") + "\n\n"
            f"[권장 조치]\n" + (action_txt or "(없음)"))


def validate_revisions(raw, sections, category_map):
    """수정 요청 검증: quote가 정확히 한 섹션 content의 '원문 그대로' 부분문자열이어야 한다.

    백엔드의 수정 승인 흐름이 anchorQuote를 exact indexOf로 찾아 교체하므로,
    공백 하나라도 다르면 여기서 탈락시킨다. 확정 정보 블록 인용도 여기서 걸러진다.
    """
    route_map = category_map["routes"]
    valid_categories = set(category_map["categories"])
    requests, dropped = [], []
    for r in (raw or {}).get("revision_requests") or []:
        quote = (r.get("quote") or "").strip()
        if not quote:
            dropped.append({"reason": "빈 인용", "raw": r})
            continue
        matches = [s.get("code") for s in sections or []
                   if s.get("code") and quote in (s.get("content") or "")]
        if len(matches) != 1:
            dropped.append({"reason": f"인용이 섹션 {len(matches)}곳과 일치", "quote": quote[:80]})
            continue
        suggestions = []
        for i, s in enumerate((r.get("suggestions") or [])[:3]):
            new_text = (s.get("new_text") or s.get("newText") or "").strip()
            if not new_text or new_text == quote:
                continue
            suggestions.append({"label": (s.get("label") or chr(65 + i)), "newText": new_text})
        if len(suggestions) < 2:
            dropped.append({"reason": f"수정안 {len(suggestions)}개(2개 미만)", "quote": quote[:80]})
            continue
        mapped = route_map.get(r.get("route_id")) or ["INDUSTRY_SPECIFIC"]
        category = next((c for c in mapped if c in valid_categories), "INDUSTRY_SPECIFIC")
        requests.append({
            "category": category,
            "anchorSectionCode": matches[0],
            "anchorQuote": quote,
            "rationale": r.get("rationale", ""),
            "suggestions": suggestions,
        })
    if dropped:
        log(f"수정 요청 검증 탈락 {len(dropped)}건: {dropped}")
    return requests


def stage_revise(state, screenings, screen_audit, sections, backend, work_dir=None):
    """고위험 근거 기반 기획서 문장 수정 요청 생성 (LLM 1회).

    실패해도 검토 전체를 죽이지 않는다 — 수정 요청 없이 결과를 돌려준다.
    """
    has_high_risk = any((screenings.get(c["citation_id"]) or {}).get("category")
                        in ("requirement", "risk")
                        for c in state.get("candidates") or [])
    if not has_high_risk:
        log("수정안 스테이지 생략: requirement/risk 근거 없음")
        return []
    prompt = build_revision_prompt(state, screenings, screen_audit)
    if work_dir:
        work_dir = Path(work_dir)
        work_dir.mkdir(parents=True, exist_ok=True)
        (work_dir / "revise_prompt.txt").write_text(
            REVISION_INSTRUCTIONS + "\n\n" + prompt, encoding="utf-8")
    try:
        raw = call_llm(prompt, REVISION_INSTRUCTIONS, backend)
    except Exception as error:  # noqa: BLE001 - 수정안은 부가 산출물, 검토 자체를 막지 않는다
        log(f"수정안 스테이지 실패(무시하고 진행): {error}")
        return []
    requests = validate_revisions(raw, sections, load_category_map())
    log(f"수정 요청 {len(requests)}건 생성")
    return requests


def review_from_sections(sections, work_dir, source_label="StructuredPlan", backend=None,
                         rerun_categories=None, confirmed_facts=None):
    """확정 계획 섹션 → 전 단계 실행. 집계에 필요한 모든 산출물을 dict로 돌려준다.

    work_dir는 요청마다 격리된 디렉터리여야 한다 (state·프롬프트가 전역 싱글턴이 되지 않도록).
    rerun_categories가 있으면 해당 범주에 매핑된 라우트만 수집·선별한다 (증분 재검토).
    confirmed_facts는 source_text에 직접 주입된다 (§4-3 확정 정보).
    """
    work_dir = Path(work_dir)
    backend = resolve_llm_backend(backend)
    text = build_source_text(sections, confirmed_facts)
    log(f"확정 계획 로드: {source_label} ({len(text):,}자)")
    _, route_prompt = stage_prepare_text(text, source_label, work_dir)

    log(f"라우팅 호출 ({MODEL} via {backend})...")
    routing_raw = call_llm(route_prompt, ROUTING_INSTRUCTIONS, backend)
    state = load_state(work_dir)
    state["llm"] = MODEL
    state["llm_backend"] = backend
    state, screen_prompt = stage_route(state, routing_raw, work_dir,
                                       rerun_categories=rerun_categories)

    log(f"선별 호출 ({MODEL} via {backend})...")
    screening_raw = call_llm(screen_prompt, SCREENING_INSTRUCTIONS, backend)
    screenings, screen_audit = validate_screening(screening_raw, state["candidates"])

    revision_requests = stage_revise(state, screenings, screen_audit, sections, backend, work_dir)
    return {"state": state, "screenings": screenings, "screen_audit": screen_audit,
            "revision_requests": revision_requests}


# ---------------------------------------------------------------- CLI

def main():
    sys.stdout.reconfigure(encoding="utf-8")
    ap = argparse.ArgumentParser(description=__doc__)
    sub = ap.add_subparsers(dest="cmd", required=True)
    p_run = sub.add_parser("run", help="전 단계 자동 실행")
    p_run.add_argument("docx")
    p_run.add_argument("--llm", default="auto",
                       choices=["auto", "anthropic", "claude-cli"],
                       help="anthropic=API 키, claude-cli=Claude Code 로그인 세션")
    p_prep = sub.add_parser("prepare-route", help="라우팅 프롬프트 생성 (수동 모드 1단계)")
    p_prep.add_argument("docx")
    p_ar = sub.add_parser("apply-route", help="라우팅 JSON 적용 → 법령 수집 → 선별 프롬프트")
    p_ar.add_argument("routing_json")
    p_as = sub.add_parser("apply-screen", help="선별 JSON 적용 → 리포트")
    p_as.add_argument("screening_json")
    args = ap.parse_args()

    if args.cmd == "run":
        backend = resolve_llm_backend(args.llm)
        text, prompt = stage_prepare(Path(args.docx))
        log(f"라우팅 호출 ({MODEL} via {backend})...")
        routing_raw = call_llm(prompt, ROUTING_INSTRUCTIONS, backend)
        state = load_state()
        state["llm"] = MODEL
        state["llm_backend"] = backend
        state, screen_prompt = stage_route(state, routing_raw)
        log(f"선별 호출 ({MODEL} via {backend})...")
        screening_raw = call_llm(screen_prompt, SCREENING_INSTRUCTIONS, backend)
        stage_screen(state, screening_raw)
    elif args.cmd == "prepare-route":
        stage_prepare(Path(args.docx))
        log("다음: route_prompt.txt를 LLM에 넣어 route_result.json 생성 후 apply-route 실행")
    elif args.cmd == "apply-route":
        state = load_state()
        routing_raw = json.loads(Path(args.routing_json).read_text(encoding="utf-8"))
        stage_route(state, routing_raw)
        log("다음: screen_prompt.txt를 LLM에 넣어 screen_result.json 생성 후 apply-screen 실행")
    elif args.cmd == "apply-screen":
        state = load_state()
        screening_raw = json.loads(Path(args.screening_json).read_text(encoding="utf-8"))
        stage_screen(state, screening_raw)


if __name__ == "__main__":
    try:
        main()
    except PipelineError as e:
        sys.exit(str(e))
