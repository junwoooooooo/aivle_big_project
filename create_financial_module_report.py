from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

OUT = Path(r"C:\dev\aivle_big_project\financial-analysis-module-report.docx")
doc = Document(); sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(.492)
styles = doc.styles
normal = styles['Normal']; normal.font.name='Calibri'; normal._element.rPr.rFonts.set(qn('w:eastAsia'),'Malgun Gothic'); normal.font.size=Pt(11); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.1
for name,size,color,before,after in [('Heading 1',16,'2E74B5',16,8),('Heading 2',13,'2E74B5',12,6),('Heading 3',12,'1F4D78',8,4)]:
 s=styles[name]; s.font.name='Calibri'; s._element.rPr.rFonts.set(qn('w:eastAsia'),'Malgun Gothic'); s.font.size=Pt(size); s.font.color.rgb=RGBColor.from_string(color); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)
header=sec.header.paragraphs[0]; header.text='AIVLE | Financial Analysis Module Design'; header.style='Normal'; header.runs[0].font.size=Pt(9); header.runs[0].font.color.rgb=RGBColor(100,100,100)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.RIGHT; footer.add_run('Internal implementation report | 2026-08-10')
def add(title, text=None):
 doc.add_heading(title, 1)
 if text: doc.add_paragraph(text)
def bullets(items):
 for x in items: doc.add_paragraph(x, style='List Bullet')
def table(headers, rows):
 t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.LEFT; t.style='Table Grid'
 for c,h in zip(t.rows[0].cells, headers): c.text=h; c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
 for row in rows:
  cells=t.add_row().cells
  for c,v in zip(cells,row): c.text=v; c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
 for row in t.rows:
  for cell in row.cells:
   tcPr=cell._tc.get_or_add_tcPr(); m=OxmlElement('w:tcMar')
   for side in ['top','start','bottom','end']:
    e=OxmlElement('w:'+side); e.set(qn('w:w'),'120' if side in ['start','end'] else '80'); e.set(qn('w:type'),'dxa'); m.append(e)
   tcPr.append(m)
 return t

p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('재무분석 모듈화 구현 보고서'); r.bold=True; r.font.size=Pt(24); r.font.color.rgb=RGBColor.from_string('0B2545'); r.font.name='Calibri'
p=doc.add_paragraph('백엔드 모듈 API · UI 샌드박스 · 시계열/몬테카를로 분석 설계', style='Subtitle'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('작성일: 2026-08-10  |  대상: AIVLE Big Project', style='Normal').alignment=WD_ALIGN_PARAGRAPH.CENTER
add('1. 구현 요약','기존 프로젝트형 재무분석 CRUD는 유지하면서, 프론트에서 독립 검증 가능한 무상태 모듈 API를 추가했다. http://localhost:3001/module 에서 금액 단위와 핵심 가정을 입력하면 API가 KRW로 변환하고 기준·보수·낙관 시나리오, 월별 현금흐름, 민감도 및 몬테카를로 위험분포를 반환한다.')
add('2. 모듈 구조와 책임'); table(['경로','역할'], [('backend/.../financial/module/FinancialModuleController.java','POST /api/v1/modules/financial/preview. UI 샌드박스 진입점.'),('.../FinancialModuleService.java','스케일링, 계산, 차트 및 보고서 DTO 조립.'),('.../FinancialInputScaler.java','입력 KRW/천원/백만원을 KRW로 단일 변환.'),('.../FinancialMonteCarloService.java','정규 충격을 적용한 반복 시뮬레이션과 분위수 계산.'),('.../module/dto/*.java','요청/응답 계약. 화면과 계산 구현의 결합 방지.'),('frontend/.../FinancialModulePage.jsx','3001/module 입력·결과 확인 화면.'),('frontend/vite.config.js','개발 포트 3001 고정.'),('backend/application.yaml','3001 origin CORS 허용.')])
add('3. 입력, DB 연계 및 스케일링'); doc.add_paragraph('현재 모듈 테스트 API는 사용자가 입력한 값으로 즉시 실행한다. 프로젝트 저장형 분석은 기존 financial_analyses 테이블의 assumptions_json, scenarios_json, result_json, summary_json 및 source_snapshot_json을 사용한다. source_snapshot_json은 프로젝트, 구조화 사업계획, 문서 버전, 타당성 분석 참조를 고정해 재현성을 보장한다.')
table(['분류','값/규칙'], [('DB에서 참조','프로젝트/사업계획/문서 버전/타당성 분석 ID 및 요약. 저장형 실행에서 source snapshot으로 고정.'),('금액','DB와 계산기는 항상 KRW. UI의 KRW·천원·백만원만 API 경계에서 1·1,000·1,000,000 배로 변환.'),('비율/수량','성장률·수수료·이탈률·판매량·구독자 수는 단위 변환하지 않는다.'),('추적성','inputHash/resultHash와 원본 snapshot으로 동일 입력 결과를 검증한다.')])
add('4. 분석 결과와 그래프','응답 calculation에는 3개 시나리오별 매출, 변동비, 고정비, 영업이익, 손익분기·회수월, 필요 운전자금이 포함된다. cashFlowChart는 월별 매출·영업이익·누적현금흐름 배열이므로 프론트에서는 선/막대 혼합 그래프로 표시한다. report는 헤드라인, 핵심 수치, 주의점, 권고 행동 및 면책문을 제공해 보고서 카드 또는 PDF/문서 생성 입력으로 쓸 수 있다.')
add('5. 시계열·민감도·몬테카를로'); bullets(['시계열 분석: 월별 판매량/구독자, 성장률, 이탈률을 반영해 12~60개월의 매출·이익·누적현금을 계산한다.', '민감도 분석: 판매량(-20~20%), 가격(-10~10%), 변동비/고정비(0~20%) 변화별 손익과 운전자금을 비교한다.', '몬테카를로 분석: 기준 가정에 판매량·가격·원가의 무작위 정규 충격을 반복 적용한다. P10/P50/P90 이익, 손실 확률, 회수 가능 확률로 단일 숫자의 불확실성을 표현한다.', '난수 seed를 응답에 보존하므로 같은 입력·seed는 같은 결과를 재현한다. 이는 AI 예측이 아니라 가정 기반 위험 시뮬레이션이다.'])
add('6. AI 연동 방침','현재 계산, 스케일링, 시계열, 몬테카를로는 결정론적 수치 엔진으로 AI 호출이 필요 없다. 향후 AI가 필요한 부분은 “수치 결과의 경영진용 서술 요약/권고”로 제한한다. ai 서비스에는 FINANCIAL_REPORT_GENERATION task type, strict Pydantic request/response 모델, prompt 파일을 추가하고 backend는 내부 실행 계약으로 요청한다. AI 출력은 제안(provenance=AI_PROPOSAL)으로만 저장하며 계산 값 자체를 변경하지 않는다.')
add('7. 저장 및 운영 권장사항'); bullets(['사용자 확정 분석은 financial_analyses에 버전·상태(DRAFT/COMPLETED)로 저장한다. 완료 결과는 수정 불가이고 복제 후 재실행한다.', '모듈 preview는 무상태이며 저장하지 않는다. “저장” 기능을 붙일 때에는 기존 프로젝트 FinancialAnalysisService create/run을 사용한다.', '대용량 시뮬레이션은 향후 analysis_jobs/task_runs로 비동기화하고 simulationCount·seed·volatility·엔진 버전을 result/source snapshot에 기록한다.', '실제 데이터가 들어오면 월별 실제 매출/비용 테이블을 별도로 두고, 가정 대비 오차와 보정 계수를 저장한다. 원본 값, 통화, 집계 기간, 데이터 출처를 함께 보관한다.'])
add('8. 검증 결과','FinancialModuleServiceTests는 백만원 입력의 KRW 정규화, 12개월 차트 배열, 기본 3개 시나리오, 동일 seed에서의 동일 몬테카를로 결과를 검증한다. 기존 FinancialCalculationServiceTests도 함께 실행해 기존 계산 계약이 유지됨을 확인했다.')
doc.save(OUT)
print(OUT)
