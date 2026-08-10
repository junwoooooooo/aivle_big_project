from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

out = Path(r"C:\dev\aivle_big_project\financial-modularization-guide.docx")
d = Document(); s = d.sections[0]
for attr in ('top_margin','bottom_margin','left_margin','right_margin'): setattr(s, attr, Inches(1))
s.header_distance=s.footer_distance=Inches(.492)
normal=d.styles['Normal']; normal.font.name='Calibri'; normal._element.rPr.rFonts.set(qn('w:eastAsia'),'Malgun Gothic'); normal.font.size=Pt(10.5); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.1
for n,size,color in [('Heading 1',16,'2E74B5'),('Heading 2',13,'2E74B5'),('Heading 3',12,'1F4D78')]:
 st=d.styles[n]; st.font.name='Calibri'; st._element.rPr.rFonts.set(qn('w:eastAsia'),'Malgun Gothic'); st.font.size=Pt(size); st.font.color.rgb=RGBColor.from_string(color); st.paragraph_format.space_before=Pt(14); st.paragraph_format.space_after=Pt(6)
h=s.header.paragraphs[0]; h.text='AIVLE | Financial Module Guide'; h.runs[0].font.size=Pt(9); h.runs[0].font.color.rgb=RGBColor(100,100,100)
f=s.footer.paragraphs[0]; f.alignment=WD_ALIGN_PARAGRAPH.RIGHT; f.add_run('Financial modularization guide | 2026-08-10')
def heading(t,l=1): d.add_heading(t,l)
def para(t): d.add_paragraph(t)
def bullets(items):
 for x in items: d.add_paragraph(x, style='List Bullet')
def table(headers, rows):
 t=d.add_table(rows=1, cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.LEFT
 for c,v in zip(t.rows[0].cells,headers): c.text=v
 for r in rows:
  cells=t.add_row().cells
  for c,v in zip(cells,r): c.text=v
 for row in t.rows:
  for cell in row.cells:
   tc=cell._tc.get_or_add_tcPr(); m=OxmlElement('w:tcMar')
   for side,val in [('top','80'),('bottom','80'),('start','120'),('end','120')]:
    e=OxmlElement('w:'+side); e.set(qn('w:w'),val); e.set(qn('w:type'),'dxa'); m.append(e)
   tc.append(m)

p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('재무분석 모듈화 가이드'); r.bold=True; r.font.name='Calibri'; r.font.size=Pt(24); r.font.color.rgb=RGBColor.from_string('0B2545')
p=d.add_paragraph('시장 · BM · 기술/운영 데이터 연계, 사용자 입력, AI 해석 및 저장 설계'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
d.add_paragraph('적용 대상: AIVLE Big Project | 2026-08-10').alignment=WD_ALIGN_PARAGRAPH.CENTER
heading('1. 목적과 원칙')
para('재무 모듈은 시장, 비즈니스 모델(BM), 기술/운영 모듈에서 검증된 근거를 받아 수익성, 현금흐름, 손익분기점(BEP), 시나리오 및 몬테카를로 위험을 계산한다. 숫자 계산은 백엔드의 결정론적 엔진이 책임지고, AI는 계산 결과를 바꾸지 않는 해석·권고 문장만 생성한다.')
bullets(['근거 우선: DB의 원본/출처/수집일을 보존하고, 누락값은 사용자 확인을 요청한다.','통화 원칙: DB와 계산기는 KRW만 사용한다. 화면의 원·천원·백만원은 API 경계에서 한 번만 KRW로 변환한다.','재현성: 입력, 시나리오, 난수 seed, 엔진 버전, AI 모델·프롬프트 버전을 저장한다.','안전성: AI 장애는 숫자 계산을 중단시키지 않으며 기본 보고서로 대체한다.'])
heading('2. 상위 모듈에서 가져올 값')
table(['원천 모듈','가져올 값','재무 반영 방식'],[
('시장','TAM/SAM/SOM, CAGR, 목표 고객 규모, 경쟁 가격, 수요 검증 결과, 출처 URL/수집일','판매량 성장률, 가격 범위, 시장 상한, 시나리오 근거와 Evidence 목록'),
('BM','수익모델(일회성/구독/혼합), 가격, 판매 채널, 목표 전환율, CAC, 재구매·이탈 가정','unitPrice, monthlySalesVolume, monthlySubscriptionPrice, initial/new subscribers, churn, marketing cost'),
('기술/운영','단위원가, 결제수수료, 인프라·서버비, 인건비, 임대료, 초기 개발/장비 투자, 생산능력','변동비·고정비·초기투자·운전자금·제약조건'),
('프로젝트/문서','프로젝트 ID, 구조화 사업계획 버전, 문서 버전, 타당성 분석 상태와 요약','source_snapshot_json에 고정하여 이후 원본 변경에도 분석 재현')])
heading('3. 사용자에게 받을 입력값')
table(['그룹','필수/조건부 값','검증'],[
('공통','분석 기간(12~60개월), 금액 단위, 수익 모델, 월 성장률, 단위 변동비, 결제수수료, 고정비, 초기투자','음수 금지; 성장률 > -100%; 수수료·이탈률 0~100%'),
('일회성 판매','제품 단가, 월 판매량','단가와 판매량은 0보다 커야 함'),
('구독','월 구독가격, 초기 구독자, 월 신규 구독자, 월 이탈률','가격은 0보다 커야 하며 구독자 수는 0 이상'),
('시나리오','보수·기준·낙관의 판매량/가격/변동비/고정비 조정률','세 코드가 모두 존재하고 범위가 현실적인지 확인'),
('시뮬레이션','반복 횟수(100~10,000), 판매량·가격·원가 변동성, 난수 seed','seed를 저장해 동일 결과를 재현')])
heading('4. 현재 백엔드 패키지 구조')
table(['폴더','핵심 파일','책임'],[
('controller','FinancialModuleController, FinancialAnalysisController, FinancialAnalysisSourceController','공개 모듈 preview와 프로젝트 저장형 CRUD API'),
('service','FinancialCalculationService, FinancialModuleService, FinancialMonteCarloService, FinancialInputScaler','수치 계산, 3개년 집계, 스트레스/몬테카를로, 단위 정규화'),
('service','FinancialAiReportClient, FinancialSourceSnapshotService, FinancialAnalysisService','AI 해석 호출, 근거 스냅샷, 버전·상태·감사 저장'),
('dto','FinancialModels, FinancialModuleRequest, FinancialModuleResponse','API/계산 전송 계약과 그래프·보고서 데이터'),
('entity/repository','FinancialAnalysis, FinancialStatus, RevenueModel, FinancialAnalysisRepository','DB 영속화와 조회')])
heading('5. 실행 흐름')
bullets(['화면 /module에서 사용자 입력 → POST /api/v1/modules/financial/preview','FinancialInputScaler가 금액을 KRW로 정규화 → FinancialCalculationService가 월별 시계열과 3개 시나리오 계산','FinancialModuleService가 3개년 손익계산서, BEP, 차트 배열, 스트레스 시나리오를 조립','FinancialMonteCarloService가 P10/P50/P90, 손실·회수 확률을 생성','FinancialAiReportClient가 AI 서버에 계산된 숫자만 보내 해석·리스크·권고를 받고, 실패 시 기본 보고서 사용','사용자 확정 시 프로젝트형 FinancialAnalysisService가 snapshot/hash/version과 함께 DB에 저장'])
heading('6. 보고서와 그래프 계약')
para('프론트는 백엔드가 내려준 annualProjections, cashFlowChart, stressScenarios, monteCarlo, report를 그대로 표시한다. 프론트에서 수익성 숫자를 다시 계산하지 않는다.')
table(['화면 구성','응답 필드'],[('구조화된 3개년 손익','revenue, variableCost, grossProfit, SG&A, operatingProfit, nonOperatingIncome, corporateTax, netIncome, margin'),('BEP 교차 그래프','cashFlowChart.month/revenue/operatingProfit/cumulativeCashFlow'),('스트레스 현금흐름','stressScenarios의 보수·기준·낙관 monthlyCashFlow'),('몬테카를로','profitP10/profitP50/profitP90/lossProbability/paybackProbability/simulations/seed'),('AI 최종 보고서','headline/findings/cautions/recommendedActions/disclaimer')])
heading('7. DB 저장 설계')
para('샌드박스 preview는 저장하지 않는다. 사용자가 저장을 확정한 프로젝트 분석만 financial_analyses에 저장한다.')
table(['컬럼','내용'],[('assumptions_json','정규화된 KRW 가정값과 수익모델'),('scenarios_json','보수·기준·낙관 조정률'),('result_json','월별·3개년 손익, BEP, 민감도, 몬테카를로, AI 보고서'),('summary_json','목록용 헤드라인과 핵심 요약'),('source_snapshot_json','시장/BM/운영 원천 ID, 문서/계획/타당성 버전, 출처'),('input_hash/result_hash','SHA-256 검증'),('version_number/status/completed_at','DRAFT→COMPLETED 상태 및 불변 결과 버전')])
heading('8. AI 및 운영 가이드')
bullets(['AI 서버 내부 엔드포인트는 /internal/v1/financial/report이며 내부 토큰으로만 호출한다.','AI 입력에는 원본 개인정보·비밀정보가 아닌 집계된 계산 결과만 포함한다.','AI 응답에는 provider, model, promptVersion, generatedAt, fallback 여부를 향후 result_json에 추가 저장한다.','시장 근거가 없으면 출처를 꾸며내지 않고 “사용자 가정 기반”으로 보고서에 표시한다.','대용량 시뮬레이션은 향후 비동기 job으로 옮기고 반복 횟수와 seed를 저장한다.'])
heading('9. 배포와 확인')
bullets(['Docker: docker compose down --remove-orphans 후 docker compose up --build -d','일반 화면은 localhost:3000, 재무 모듈은 localhost:3001/module','backend와 ai-server의 AI_INTERNAL_SERVICE_TOKEN / AI_SERVER_INTERNAL_API_KEY는 동일한 내부 토큰을 사용한다.','/jaemu 및 별도 jaemu 컨테이너는 제거되었으며 재무 기능은 financial 패키지에만 존재한다.'])
d.save(out); print(out)
